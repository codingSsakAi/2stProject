import os
import re
import easyocr
from docx import Document
from docx.shared import Pt

# PNG 이미지 폴더
img_dir = 'pdf_png'
img_files = sorted(
    [f for f in os.listdir(img_dir) if f.lower().endswith('.png')],
    key=lambda x: int(re.findall(r'\d+', x)[0])  # page_0.png, page_1.png 등 번호순 정렬
)

reader = easyocr.Reader(['ko', 'en'])
doc = Document()
doc.styles['Normal'].font.name = '맑은 고딕'
doc.styles['Normal'].font.size = Pt(11)

for i, img_file in enumerate(img_files):
    img_path = os.path.join(img_dir, img_file)
    ocr_result = reader.readtext(img_path, detail=0, paragraph=True)
    text = "\n".join(ocr_result).strip()
    doc.add_paragraph(f"[{i+1}페이지]")
    doc.add_paragraph(text)
    doc.add_paragraph("\n" + "-"*40 + "\n")

docx_path = 'ocr_result_from_png.docx'
doc.save(docx_path)
print(f"변환 완료: {docx_path}")
