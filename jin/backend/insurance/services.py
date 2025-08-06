"""
RAG 시스템 서비스
문서 검색 및 응답 생성을 위한 RAG 시스템
"""

import logging
import os
from typing import List, Dict, Any, Optional
from django.conf import settings
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
import pinecone
from openai import OpenAI
from langchain_upstage import UpstageEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Pinecone
from langchain.chains import RetrievalQA
from langchain_community.llms import OpenAI as LangChainOpenAI
import PyPDF2
from docx import Document
import io

# 로깅 설정
logger = logging.getLogger(__name__)


class RAGService:
    """RAG 시스템 서비스 클래스"""

    def __init__(self):
        """RAG 서비스 초기화"""
        self._initialize_pinecone()
        self._initialize_openai()
        self._initialize_embeddings()
        self._initialize_vectorstore()

    def _initialize_pinecone(self):
        """Pinecone 초기화"""
        try:
            api_key = settings.PINECONE_API_KEY
            environment = settings.PINECONE_ENVIRONMENT
            index_name = settings.PINECONE_INDEX_NAME

            if not api_key:
                logger.warning("Pinecone API 키가 설정되지 않았습니다.")
                self.pinecone_index = None
                return

            # 여러 환경 시도
            environments_to_try = [
                environment,
                "us-east-1-aws",
                "us-west1-gcp",
                "gcp-starter",
            ]

            pinecone_initialized = False
            for env in environments_to_try:
                try:
                    logger.info(f"Pinecone 환경 시도: {env}")
                    pinecone.init(api_key=api_key, environment=env)

                    # 인덱스 목록 확인
                    available_indexes = pinecone.list_indexes()
                    logger.info(f"사용 가능한 Pinecone 인덱스: {available_indexes}")

                    # 인덱스가 존재하지 않으면 생성
                    if index_name not in available_indexes:
                        try:
                            pinecone.create_index(
                                name=index_name,
                                dimension=settings.UPSTAGE_EMBEDDING_DIMENSION,
                                metric="cosine",
                            )
                            logger.info(f"Pinecone 인덱스 생성: {index_name}")
                        except Exception as create_error:
                            logger.error(f"Pinecone 인덱스 생성 실패: {create_error}")
                            continue

                    self.pinecone_index = pinecone.Index(index_name)
                    logger.info(f"Pinecone 초기화 완료 (환경: {env})")
                    pinecone_initialized = True
                    break

                except Exception as e:
                    logger.warning(f"Pinecone 환경 {env} 연결 실패: {e}")
                    continue

            if not pinecone_initialized:
                logger.error("모든 Pinecone 환경 연결 실패")
                self.pinecone_index = None

        except Exception as e:
            logger.error(f"Pinecone 초기화 실패: {e}")
            self.pinecone_index = None

    def _initialize_openai(self):
        """OpenAI 초기화"""
        try:
            api_key = settings.OPENAI_API_KEY
            if not api_key:
                logger.warning("OpenAI API 키가 설정되지 않았습니다.")
                self.openai_client = None
                return

            self.openai_client = OpenAI(api_key=api_key)
            logger.info("OpenAI 초기화 완료")

        except Exception as e:
            logger.error(f"OpenAI 초기화 실패: {e}")
            self.openai_client = None

    def _initialize_embeddings(self):
        """임베딩 모델 초기화"""
        try:
            api_key = settings.UPSTAGE_API_KEY
            model_name = settings.UPSTAGE_EMBEDDING_MODEL

            if not api_key:
                logger.warning("Upstage API 키가 설정되지 않았습니다.")
                self.embeddings = None
                return

            self.embeddings = UpstageEmbeddings(model=model_name, api_key=api_key)
            logger.info("Upstage 임베딩 초기화 완료")

        except Exception as e:
            logger.error(f"임베딩 초기화 실패: {e}")
            self.embeddings = None

    def _initialize_vectorstore(self):
        """벡터 스토어 초기화"""
        try:
            if self.pinecone_index is None or self.embeddings is None:
                logger.warning(
                    "Pinecone 또는 임베딩이 초기화되지 않아 벡터 스토어를 생성할 수 없습니다."
                )
                self.vectorstore = None
                return

            # 빈 인덱스로 시작할 수 있도록 수정
            self.vectorstore = Pinecone(
                index=self.pinecone_index,
                embedding=self.embeddings,
                text_key="text"
            )
            logger.info("벡터 스토어 초기화 완료")

        except Exception as e:
            logger.error(f"벡터 스토어 초기화 실패: {e}")
            self.vectorstore = None

    def _extract_text_from_pdf(self, file) -> str:
        """PDF 파일에서 텍스트 추출"""
        try:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
        except Exception as e:
            logger.error(f"PDF 텍스트 추출 실패: {e}")
            return ""

    def _extract_text_from_docx(self, file) -> str:
        """DOCX 파일에서 텍스트 추출"""
        try:
            doc = Document(file)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            logger.error(f"DOCX 텍스트 추출 실패: {e}")
            return ""

    def _split_text(self, text: str) -> List[str]:
        """텍스트를 청크로 분할"""
        try:
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=settings.CHUNK_SIZE,
                chunk_overlap=settings.CHUNK_OVERLAP,
                length_function=len,
            )
            chunks = splitter.split_text(text)
            return chunks
        except Exception as e:
            logger.error(f"텍스트 분할 실패: {e}")
            return [text]

    def upload_document(self, file) -> Dict[str, Any]:
        """문서 업로드 및 벡터화"""
        try:
            if self.vectorstore is None:
                return {
                    "success": False,
                    "error": "벡터 스토어가 초기화되지 않았습니다.",
                }

            # 파일 확장자 확인
            file_extension = file.name.split(".")[-1].lower()

            # 텍스트 추출
            if file_extension == "pdf":
                text = self._extract_text_from_pdf(file)
            elif file_extension == "docx":
                text = self._extract_text_from_docx(file)
            else:
                return {"success": False, "error": "지원하지 않는 파일 형식입니다."}

            if not text.strip():
                return {
                    "success": False,
                    "error": "파일에서 텍스트를 추출할 수 없습니다.",
                }

            # 텍스트 분할
            chunks = self._split_text(text)

            # 벡터 스토어에 추가
            metadatas = [{"source": file.name, "chunk": i} for i in range(len(chunks))]
            self.vectorstore.add_texts(texts=chunks, metadatas=metadatas)

            return {
                "success": True,
                "filename": file.name,
                "chunks_count": len(chunks),
                "text_length": len(text),
            }

        except Exception as e:
            logger.error(f"문서 업로드 실패: {e}")
            return {"success": False, "error": str(e)}

    def upload_document_with_metadata(self, file_path, company, document_type, title, tags=""):
        """보험사별 메타데이터를 포함한 문서 업로드"""
        try:
            if self.vectorstore is None:
                return {
                    "success": False,
                    "error": "벡터 스토어가 초기화되지 않았습니다.",
                }

            # 파일 확장자 확인
            file_extension = str(file_path).split(".")[-1].lower()

            # 텍스트 추출
            if file_extension == "pdf":
                text = self._extract_text_from_pdf_file(file_path)
            elif file_extension == "docx":
                text = self._extract_text_from_docx_file(file_path)
            else:
                return {"success": False, "error": "지원하지 않는 파일 형식입니다."}

            if not text.strip():
                return {
                    "success": False,
                    "error": "파일에서 텍스트를 추출할 수 없습니다.",
                }

            # 텍스트 분할
            chunks = self._split_text(text)

            # 보험사별 메타데이터 생성
            from datetime import datetime
            upload_date = datetime.now().strftime("%Y-%m-%d")
            
            metadatas = []
            for i in range(len(chunks)):
                metadata = {
                    "source": str(file_path).split("/")[-1],
                    "company": company,
                    "document_type": document_type,
                    "title": title,
                    "upload_date": upload_date,
                    "chunk": i,
                    "file_path": str(file_path),
                    "tags": tags,
                    "total_chunks": len(chunks)
                }
                metadatas.append(metadata)

            # 벡터 스토어에 추가
            self.vectorstore.add_texts(texts=chunks, metadatas=metadatas)

            return {
                "success": True,
                "filename": str(file_path).split("/")[-1],
                "company": company,
                "document_type": document_type,
                "chunks_count": len(chunks),
                "text_length": len(text),
                "upload_date": upload_date
            }

        except Exception as e:
            logger.error(f"문서 업로드 실패: {e}")
            return {"success": False, "error": str(e)}

    def _extract_text_from_pdf_file(self, file_path) -> str:
        """PDF 파일 경로에서 텍스트 추출"""
        try:
            pdf_reader = PyPDF2.PdfReader(file_path)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
        except Exception as e:
            logger.error(f"PDF 텍스트 추출 실패: {e}")
            return ""

    def _extract_text_from_docx_file(self, file_path) -> str:
        """DOCX 파일 경로에서 텍스트 추출"""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            logger.error(f"DOCX 텍스트 추출 실패: {e}")
            return ""

    def search_documents(self, query: str, top_k: int = 5) -> List[Dict[str, Any]]:
        """문서 검색"""
        try:
            if self.vectorstore is None:
                logger.warning(
                    "벡터 스토어가 초기화되지 않아 검색을 수행할 수 없습니다."
                )
                return []

            # 유사도 검색
            docs = self.vectorstore.similarity_search(query, k=top_k)

            results = []
            for doc in docs:
                results.append(
                    {
                        "content": doc.page_content,
                        "metadata": doc.metadata,
                        "score": getattr(doc, "score", 0.0),
                    }
                )

            return results

        except Exception as e:
            logger.error(f"문서 검색 실패: {e}")
            return []

    def search_documents_by_company(self, query: str, company: str = None, top_k: int = 5) -> List[Dict[str, Any]]:
        """보험사별 문서 검색"""
        try:
            if self.vectorstore is None:
                logger.warning("벡터 스토어가 초기화되지 않아 검색을 수행할 수 없습니다.")
                return []

            # 필터 설정
            filter_dict = {}
            if company:
                filter_dict["company"] = company

            # 유사도 검색
            if filter_dict:
                docs = self.vectorstore.similarity_search(query, k=top_k, filter=filter_dict)
            else:
                docs = self.vectorstore.similarity_search(query, k=top_k)

            results = []
            for doc in docs:
                results.append({
                    "content": doc.page_content,
                    "metadata": doc.metadata,
                    "score": getattr(doc, "score", 0.0),
                })

            return results

        except Exception as e:
            logger.error(f"보험사별 문서 검색 실패: {e}")
            return []

    def generate_response(self, query: str) -> str:
        """응답 생성"""
        try:
            if self.openai_client is None:
                return "OpenAI 서비스가 초기화되지 않았습니다."

            # 관련 문서 검색
            relevant_docs = self.search_documents(query, top_k=3)

            if not relevant_docs:
                return "관련 문서를 찾을 수 없습니다."

            # 컨텍스트 구성
            context = "\n\n".join([doc["content"] for doc in relevant_docs])

            # 프롬프트 구성
            prompt = f"""
            다음 컨텍스트를 바탕으로 질문에 답변해주세요.
            
            컨텍스트:
            {context}
            
            질문: {query}
            
            답변:
            """

            # OpenAI API 호출
            response = self.openai_client.chat.completions.create(
                model=settings.OPENAI_MODEL,
                messages=[
                    {
                        "role": "system",
                        "content": "당신은 자동차 보험 전문 상담사입니다. 정확하고 도움이 되는 답변을 제공해주세요.",
                    },
                    {"role": "user", "content": prompt},
                ],
                max_tokens=settings.OPENAI_MAX_TOKENS,
                temperature=settings.OPENAI_TEMPERATURE,
            )

            return response.choices[0].message.content

        except Exception as e:
            logger.error(f"응답 생성 실패: {e}")
            return f"응답 생성 중 오류가 발생했습니다: {str(e)}"

    def chat(self, message: str) -> str:
        """채팅 응답"""
        return self.generate_response(message)

    def delete_document(self, document_id: str) -> Dict[str, Any]:
        """문서 삭제"""
        try:
            if self.pinecone_index is None:
                return {"success": False, "error": "Pinecone이 초기화되지 않았습니다."}

            # Pinecone에서 문서 삭제
            self.pinecone_index.delete(ids=[document_id])

            return {"success": True, "message": "문서가 삭제되었습니다."}

        except Exception as e:
            logger.error(f"문서 삭제 실패: {e}")
            return {"success": False, "error": str(e)}

    def get_index_stats(self) -> Dict[str, Any]:
        """인덱스 통계 조회"""
        try:
            if self.pinecone_index is None:
                return {
                    "total_documents": 0,
                    "total_companies": 0,
                    "index_size": 0,
                    "embedding_dimension": settings.UPSTAGE_EMBEDDING_DIMENSION,
                }

            # Pinecone 인덱스 통계
            index_stats = self.pinecone_index.describe_index_stats()

            return {
                "total_documents": index_stats.get("total_vector_count", 0),
                "total_companies": 0,  # TODO: 보험사 수 계산
                "index_size": index_stats.get("dimension", 0),
                "embedding_dimension": settings.UPSTAGE_EMBEDDING_DIMENSION,
            }

        except Exception as e:
            logger.error(f"인덱스 통계 조회 실패: {e}")
            return {
                "total_documents": 0,
                "total_companies": 0,
                "index_size": 0,
                "embedding_dimension": settings.UPSTAGE_EMBEDDING_DIMENSION,
            }

    def get_company_document_stats(self) -> Dict[str, Any]:
        """보험사별 문서 통계"""
        try:
            if self.vectorstore is None:
                return {"error": "벡터 스토어가 초기화되지 않았습니다."}

            # Pinecone에서 모든 벡터 조회
            stats = self.pinecone_index.describe_index_stats()
            
            # 보험사별 통계 계산
            company_stats = {}
            total_vectors = stats.get("total_vector_count", 0)
            
            # 보험사 목록
            companies = [
                "삼성화재", "현대해상", "KB손해보험", "메리츠화재",
                "DB손해보험", "롯데손해보험", "하나손해보험", "흥국화재",
                "AXA손해보험", "MG손해보험", "캐롯손해보험", "한화손해보험"
            ]
            
            for company in companies:
                try:
                    # 보험사별 벡터 수 조회
                    company_vectors = self.pinecone_index.query(
                        vector=[0] * settings.UPSTAGE_EMBEDDING_DIMENSION,
                        top_k=1,
                        filter={"company": company},
                        include_metadata=False
                    )
                    company_stats[company] = len(company_vectors.matches)
                except Exception as e:
                    company_stats[company] = 0
                    logger.warning(f"{company} 통계 조회 실패: {e}")

            return {
                "total_vectors": total_vectors,
                "company_stats": company_stats,
                "index_dimension": settings.UPSTAGE_EMBEDDING_DIMENSION,
                "index_name": settings.PINECONE_INDEX_NAME
            }

        except Exception as e:
            logger.error(f"보험사별 문서 통계 조회 실패: {e}")
            return {"error": str(e)}
