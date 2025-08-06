"""
Insurance 앱 뷰
RAG 시스템 및 보험 추천 시스템을 위한 뷰
"""

import logging
import json
from typing import Dict, Any, List
from django.shortcuts import render
from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods
from django.conf import settings
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import AllowAny
from rest_framework.response import Response
from rest_framework import status

from .services import RAGService
from .models import PolicyDocument, InsuranceCompany

from django.contrib.auth.decorators import user_passes_test
import os
from pathlib import Path
from django.conf import settings
import PyPDF2
from docx import Document
import logging
import re

# 로깅 설정
logger = logging.getLogger(__name__)


def main_page(request):
    """메인 페이지 - AI 챗봇 중심 보험 추천 시스템 홈"""
    try:
        context = {
            "title": "AI 챗봇 자동차 보험 추천",
            "description": "개인화된 AI 챗봇으로 최적의 자동차 보험을 찾아보세요",
        }

        return render(request, "insurance/main_page.jinja.html", context)

    except Exception as e:
        logger.error(f"메인 페이지 로드 실패: {e}")
        return render(
            request,
            "insurance/main_page.jinja.html",
            {"error": f"페이지 로드 중 오류가 발생했습니다: {str(e)}"},
        )


def rag_dashboard(request):
    """RAG 대시보드"""
    try:
        rag_service = RAGService()
        stats = rag_service.get_index_stats()

        context = {"title": "RAG 시스템 대시보드", "stats": stats}

        return render(request, "insurance/rag_dashboard.jinja.html", context)

    except Exception as e:
        logger.error(f"RAG 대시보드 로드 실패: {e}")
        return HttpResponse(
            f"대시보드 로드 중 오류가 발생했습니다: {str(e)}", status=500
        )


def compare_insurance(request):
    """보험 비교 페이지"""
    try:
        context = {
            "title": "보험 상품 비교",
            "description": "다양한 자동차 보험 상품을 비교해보세요",
        }

        return render(request, "insurance/compare.jinja.html", context)

    except Exception as e:
        logger.error(f"보험 비교 페이지 로드 실패: {e}")
        return render(
            request,
            "insurance/compare.jinja.html",
            {"error": f"페이지 로드 중 오류가 발생했습니다: {str(e)}"},
        )


def about_page(request):
    """소개 페이지"""
    try:
        context = {
            "title": "서비스 소개",
            "description": "AI 기반 자동차 보험 추천 시스템에 대해 알아보세요",
        }

        return render(request, "insurance/about.jinja.html", context)

    except Exception as e:
        logger.error(f"소개 페이지 로드 실패: {e}")
        return render(
            request,
            "insurance/about.jinja.html",
            {"error": f"페이지 로드 중 오류가 발생했습니다: {str(e)}"},
        )


def personalized_chat(request):
    """개인화된 AI 챗봇 페이지"""
    try:
        context = {
            "title": "개인화된 AI 챗봇",
            "description": "당신의 프로필을 기반으로 맞춤형 보험 상담을 제공합니다",
        }

        return render(request, "insurance/personalized_chat.jinja.html", context)

    except Exception as e:
        logger.error(f"개인화 챗봇 페이지 로드 실패: {e}")
        return render(
            request,
            "insurance/personalized_chat.jinja.html",
            {"error": f"페이지 로드 중 오류가 발생했습니다: {str(e)}"},
        )


@api_view(["POST"])
@permission_classes([AllowAny])
def search_documents_api(request):
    """문서 검색 API"""
    try:
        data = request.data
        query = data.get("query", "").strip()

        if not query:
            return Response(
                {"error": "검색어가 비어있습니다."}, status=status.HTTP_400_BAD_REQUEST
            )

        rag_service = RAGService()
        results = rag_service.search_documents(query, top_k=5)

        return Response({"success": True, "query": query, "results": results})

    except Exception as e:
        logger.error(f"문서 검색 실패: {e}")
        return Response(
            {"error": f"검색 중 오류가 발생했습니다: {str(e)}"},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR,
        )


@api_view(["POST"])
@permission_classes([AllowAny])
def generate_response_api(request):
    """응답 생성 API"""
    try:
        data = request.data
        query = data.get("query", "").strip()

        if not query:
            return Response(
                {"error": "질문이 비어있습니다."}, status=status.HTTP_400_BAD_REQUEST
            )

        rag_service = RAGService()
        response = rag_service.generate_response(query)

        return Response({"success": True, "query": query, "response": response})

    except Exception as e:
        logger.error(f"응답 생성 실패: {e}")
        return Response(
            {"error": f"응답 생성 중 오류가 발생했습니다: {str(e)}"},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR,
        )


@api_view(["POST"])
@permission_classes([AllowAny])
def chat_api(request):
    """채팅 API"""
    try:
        data = request.data
        message = data.get("message", "").strip()

        if not message:
            return Response(
                {"error": "메시지가 비어있습니다."}, status=status.HTTP_400_BAD_REQUEST
            )

        rag_service = RAGService()
        response = rag_service.chat(message)

        return Response({"success": True, "message": message, "response": response})

    except Exception as e:
        logger.error(f"채팅 실패: {e}")
        return Response(
            {"error": f"채팅 중 오류가 발생했습니다: {str(e)}"},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR,
        )


@api_view(["GET"])
@permission_classes([AllowAny])
def get_index_stats_api(request):
    """인덱스 통계 API"""
    try:
        rag_service = RAGService()
        stats = rag_service.get_index_stats()

        return Response({"success": True, "stats": stats})

    except Exception as e:
        logger.error(f"통계 조회 실패: {e}")
        return Response(
            {"error": f"통계 조회 중 오류가 발생했습니다: {str(e)}"},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR,
        )


@api_view(["POST"])
@permission_classes([AllowAny])
def upload_document_api(request):
    """문서 업로드 API"""
    try:
        if "file" not in request.FILES:
            return Response(
                {"error": "파일이 업로드되지 않았습니다."},
                status=status.HTTP_400_BAD_REQUEST,
            )

        file = request.FILES["file"]

        # 파일 크기 검증
        if file.size > settings.MAX_UPLOAD_SIZE:
            return Response(
                {
                    "error": f"파일 크기가 너무 큽니다. 최대 {settings.MAX_UPLOAD_SIZE} bytes"
                },
                status=status.HTTP_400_BAD_REQUEST,
            )

        # 파일 확장자 검증
        file_extension = file.name.split(".")[-1].lower()
        if file_extension not in settings.ALLOWED_FILE_TYPES:
            return Response(
                {
                    "error": f'지원하지 않는 파일 형식입니다. 지원 형식: {", ".join(settings.ALLOWED_FILE_TYPES)}'
                },
                status=status.HTTP_400_BAD_REQUEST,
            )

        rag_service = RAGService()
        result = rag_service.upload_document(file)

        return Response(
            {
                "success": True,
                "message": "문서가 성공적으로 업로드되었습니다.",
                "result": result,
            }
        )

    except Exception as e:
        logger.error(f"문서 업로드 실패: {e}")
        return Response(
            {"error": f"문서 업로드 중 오류가 발생했습니다: {str(e)}"},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR,
        )


@api_view(["POST"])
@permission_classes([AllowAny])
def delete_document_api(request):
    """문서 삭제 API"""
    try:
        data = request.data
        document_id = data.get("document_id")

        if not document_id:
            return Response(
                {"error": "문서 ID가 필요합니다."}, status=status.HTTP_400_BAD_REQUEST
            )

        rag_service = RAGService()
        result = rag_service.delete_document(document_id)

        return Response(
            {
                "success": True,
                "message": "문서가 성공적으로 삭제되었습니다.",
                "result": result,
            }
        )

    except Exception as e:
        logger.error(f"문서 삭제 실패: {e}")
        return Response(
            {"error": f"문서 삭제 중 오류가 발생했습니다: {str(e)}"},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR,
        )


def admin_required(view_func):
    """관리자 권한 확인 데코레이터"""
    return user_passes_test(lambda u: u.is_staff)(view_func)


@admin_required
def admin_upload_document(request):
    """관리자 문서 업로드 페이지 (단순화된 버전)"""
    if request.method == "POST":
        try:
            company = request.POST.get("company")
            document_file = request.FILES.get("document")

            if not company or not document_file:
                return JsonResponse(
                    {"success": False, "error": "보험사와 파일을 선택해주세요."}
                )

            # 자동으로 제목과 타입 생성
            filename = secure_filename(document_file.name)
            title = filename.replace(".pdf", "").replace(".docx", "")
            document_type = "이용약관"  # 기본값

            # 파일 저장 및 처리
            result = process_document_upload(
                company, document_file, title, document_type, "", ""
            )

            return JsonResponse(result)

        except Exception as e:
            return JsonResponse({"success": False, "error": str(e)})

    # GET 요청: 업로드 페이지 표시
    insurance_companies = [
        "삼성화재",
        "현대해상",
        "KB손해보험",
        "메리츠화재",
        "DB손해보험",
        "롯데손해보험",
        "하나손해보험",
        "흥국화재",
        "AXA손해보험",
        "MG손해보험",
        "캐롯손해보험",
        "한화손해보험",
    ]

    context = {"insurance_companies": insurance_companies, "title": "단일 파일 업로드"}

    return render(request, "insurance/admin_upload.jinja.html", context)


def process_document_upload(company, file, title, document_type, description, tags):
    """문서 업로드 처리"""
    try:
        # 1. PDF 폴더에 파일 저장
        pdf_dir = Path(settings.BASE_DIR) / "policy_documents" / "pdf" / company
        pdf_dir.mkdir(parents=True, exist_ok=True)

        filename = secure_filename(file.name)
        pdf_path = pdf_dir / filename

        with open(pdf_path, "wb") as f:
            for chunk in file.chunks():
                f.write(chunk)

        # 2. PDF → DOCX 변환
        docx_dir = Path(settings.BASE_DIR) / "policy_documents" / "docx" / company
        docx_dir.mkdir(parents=True, exist_ok=True)

        docx_filename = filename.replace(".pdf", ".docx")
        docx_path = docx_dir / docx_filename

        # PDF → DOCX 변환
        convert_pdf_to_docx(pdf_path, docx_path)

        # 3. 데이터베이스에 문서 정보 저장
        document = PolicyDocument.objects.create(
            title=title,
            company=InsuranceCompany.objects.get(name=company),
            document_type=document_type,
            description=description,
            tags=tags,
            file_path=str(docx_path),
            status="approved",
        )

        # 4. Pinecone에 업로드
        rag_service = RAGService()
        upload_result = rag_service.upload_document_with_metadata(
            docx_path, company, document_type, title, tags
        )

        return {
            "success": True,
            "message": f"문서 업로드 완료: {title}",
            "document_id": document.id,
            "pinecone_result": upload_result,
        }

    except Exception as e:
        logger.error(f"문서 업로드 실패: {e}")
        return {"success": False, "error": str(e)}


def convert_pdf_to_docx(pdf_path, docx_path):
    """PDF를 DOCX로 변환 (개선된 버전)"""
    try:
        # 방법 1: PyPDF2 사용
        text = extract_text_with_pypdf2(pdf_path)

        # 방법 1이 실패하면 방법 2 시도
        if not text.strip():
            text = extract_text_with_alternative_method(pdf_path)

        if not text.strip():
            raise Exception("PDF에서 텍스트를 추출할 수 없습니다.")

        # DOCX 생성
        doc = Document()

        # 텍스트를 문단으로 분할하여 추가
        paragraphs = text.split("\n")
        for para in paragraphs:
            if para.strip():
                # 특수 문자 제거 후 문단 추가
                clean_para = clean_text_for_docx(para)
                if clean_para.strip():
                    doc.add_paragraph(clean_para)

        doc.save(docx_path)
        logger.info(f"PDF → DOCX 변환 완료: {pdf_path} → {docx_path}")

    except Exception as e:
        logger.error(f"PDF → DOCX 변환 실패: {e}")
        raise e


def extract_text_with_pypdf2(pdf_path):
    """PyPDF2를 사용한 텍스트 추출 (한글 지원 개선)"""
    try:
        import PyPDF2

        with open(pdf_path, "rb") as file:
            pdf_reader = PyPDF2.PdfReader(file)

            text = ""
            for page_num in range(len(pdf_reader.pages)):
                try:
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()

                    if page_text:
                        # 한글 텍스트 정리
                        cleaned_text = clean_korean_text(page_text)
                        if cleaned_text.strip():
                            text += cleaned_text + "\n\n"

                except Exception as e:
                    logger.warning(f"페이지 {page_num + 1} 텍스트 추출 실패: {e}")
                    continue

            return text.strip()

    except Exception as e:
        logger.error(f"PyPDF2 텍스트 추출 실패: {e}")
        return ""


def clean_korean_text(text):
    """한글 텍스트 정리"""
    import re

    # 기본 정리
    text = text.replace("\x00", "")  # NULL 바이트 제거

    # 한글 문자 정규화
    text = re.sub(
        r"[^\uAC00-\uD7AF\u1100-\u11FF\u3130-\u318F\uA960-\uA97F\uAC00-\uD7AF\uD7B0-\uD7FF\w\s\.\,\;\:\!\?\(\)\[\]\-\_\+\=\*\&\^\%\$\#\@\~]",
        "",
        text,
    )

    # 연속된 공백 정리
    text = re.sub(r"\s+", " ", text)

    # 빈 줄 정리
    text = re.sub(r"\n\s*\n", "\n", text)

    return text.strip()


def extract_text_with_alternative_method(pdf_path):
    """대체 방법을 사용한 텍스트 추출 (한글 지원 강화)"""
    try:
        # 방법 1: pdfplumber 시도 (한글 지원 우수)
        try:
            import pdfplumber

            with pdfplumber.open(pdf_path) as pdf:
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        cleaned_text = clean_korean_text(page_text)
                        if cleaned_text.strip():
                            text += cleaned_text + "\n\n"

                if text.strip():
                    return clean_text_for_docx(text)
        except ImportError:
            logger.info("pdfplumber가 설치되지 않음, 다른 방법 시도")
        except Exception as e:
            logger.warning(f"pdfplumber 추출 실패: {e}")

        # 방법 2: PyMuPDF (fitz) 시도
        try:
            import fitz

            doc = fitz.open(pdf_path)
            text = ""
            for page in doc:
                page_text = page.get_text()
                if page_text:
                    cleaned_text = clean_korean_text(page_text)
                    if cleaned_text.strip():
                        text += cleaned_text + "\n\n"
            doc.close()

            if text.strip():
                return clean_text_for_docx(text)
        except ImportError:
            logger.info("PyMuPDF가 설치되지 않음, 다른 방법 시도")
        except Exception as e:
            logger.warning(f"PyMuPDF 추출 실패: {e}")

        # 방법 3: 기존 방법 (개선됨)
        with open(pdf_path, "rb") as file:
            content = file.read()

        if not content.startswith(b"%PDF"):
            return ""

        text = ""

        # 다양한 인코딩 시도
        encodings = ["utf-8", "cp949", "euc-kr", "latin-1", "iso-8859-1"]

        for encoding in encodings:
            try:
                content_str = content.decode(encoding, errors="ignore")
                extracted_text = extract_text_from_string_improved(content_str)
                if extracted_text.strip():
                    text = extracted_text
                    break
            except Exception as e:
                logger.debug(f"{encoding} 디코딩 실패: {e}")
                continue

        return clean_text_for_docx(text)

    except Exception as e:
        logger.warning(f"대체 방법 추출 실패: {e}")
        return ""


def extract_text_from_string_improved(content_str):
    """개선된 문자열에서 텍스트 추출 (한글 지원)"""
    import re

    text = ""

    # 한글 텍스트 패턴들
    korean_patterns = [
        r"([가-힣\s]{10,})",  # 한글 + 공백 (10자 이상)
        r"([가-힣\w\s]{15,})",  # 한글 + 영문 + 공백 (15자 이상)
        r"\(([가-힣\w\s]{5,})\)",  # 괄호 안의 한글 텍스트
        r"\[([가-힣\w\s]{5,})\]",  # 대괄호 안의 한글 텍스트
    ]

    # PDF 텍스트 명령어 패턴들
    pdf_patterns = [
        r"BT\s*([^E]*?)\s*ET",  # 텍스트 블록
        r"Tj\s*([^)]*)\)",  # Tj 명령어
        r"TJ\s*\[([^\]]*)\]",  # TJ 명령어
    ]

    # 한글 패턴 먼저 시도
    for pattern in korean_patterns:
        try:
            matches = re.findall(pattern, content_str, re.DOTALL)
            for match in matches:
                if isinstance(match, str) and len(match.strip()) > 5:
                    # 한글 비율 확인
                    korean_chars = len(re.findall(r"[가-힣]", match))
                    if korean_chars > 0:  # 한글이 포함된 경우만
                        clean_match = re.sub(r"[^\w\s가-힣]", "", match)
                        if len(clean_match) > 5:
                            text += clean_match.strip() + "\n"
        except Exception as e:
            logger.debug(f"한글 패턴 {pattern} 처리 실패: {e}")
            continue

    # PDF 패턴 시도
    for pattern in pdf_patterns:
        try:
            matches = re.findall(pattern, content_str, re.DOTALL)
            for match in matches:
                if isinstance(match, str) and len(match.strip()) > 5:
                    # 한글이 포함된 경우만 선택
                    if re.search(r"[가-힣]", match):
                        clean_match = re.sub(r"[^\w\s가-힣]", "", match)
                        if len(clean_match) > 5:
                            text += clean_match.strip() + "\n"
        except Exception as e:
            logger.debug(f"PDF 패턴 {pattern} 처리 실패: {e}")
            continue

    return text


def clean_text_for_docx(text):
    """DOCX 호환을 위한 텍스트 정리"""
    import re

    # NULL 바이트 제거
    text = text.replace("\x00", "")

    # 제어 문자 제거 (0x00-0x1F, 0x7F-0x9F)
    text = "".join(char for char in text if ord(char) >= 32 or char in "\n\r\t")

    # XML 호환되지 않는 문자 제거
    text = re.sub(
        r"[^\x20-\x7E\xA0-\xFF\u0100-\u017F\u0180-\u024F\u1E00-\u1EFF\u2C60-\u2C7F\uA720-\uA7FF\uFB00-\uFB4F]",
        "",
        text,
    )

    # 연속된 공백 정리
    text = re.sub(r"\s+", " ", text)

    # 빈 줄 정리
    text = re.sub(r"\n\s*\n", "\n", text)

    return text.strip()


def convert_multiple_pdfs_to_single_docx(pdf_files, company):
    """여러 PDF 파일을 하나의 DOCX로 통합 변환"""
    try:
        from docx import Document
        from docx.shared import Inches
        from pathlib import Path

        # 새로운 DOCX 문서 생성
        doc = Document()

        # 제목 추가
        title = doc.add_heading(f"{company} 이용약관 통합본", 0)
        title.alignment = 1  # 가운데 정렬

        # 날짜 추가
        from datetime import datetime

        date_paragraph = doc.add_paragraph(
            f"생성일: {datetime.now().strftime('%Y년 %m월 %d일')}"
        )
        date_paragraph.alignment = 1  # 가운데 정렬

        # 구분선 추가
        doc.add_paragraph("=" * 50)

        all_text = ""
        file_count = 0

        for pdf_file in pdf_files:
            try:
                # pdf_file이 문자열인지 파일 객체인지 확인
                if isinstance(pdf_file, str):
                    pdf_path = pdf_file
                    filename = Path(pdf_file).name
                else:
                    pdf_path = str(pdf_file)
                    filename = (
                        pdf_file.name
                        if hasattr(pdf_file, "name")
                        else Path(pdf_file).name
                    )

                # PDF에서 텍스트 추출
                text = extract_text_with_pypdf2(pdf_path)

                if not text.strip():
                    text = extract_text_with_alternative_method(pdf_path)

                if text.strip():
                    # 파일명을 제목으로 추가
                    clean_filename = filename.replace(".pdf", "")
                    doc.add_heading(f"【{clean_filename}】", level=1)

                    # 텍스트를 문단으로 분할하여 추가
                    paragraphs = text.split("\n")
                    for para in paragraphs:
                        if para.strip():
                            clean_para = clean_text_for_docx(para)
                            if clean_para.strip():
                                doc.add_paragraph(clean_para)

                    # 구분선 추가
                    doc.add_paragraph("-" * 30)
                    file_count += 1
                    all_text += text + "\n\n"

            except Exception as e:
                logger.error(f"PDF 처리 실패: {filename} - {e}")
                continue

        if file_count == 0:
            raise Exception("처리할 수 있는 PDF 파일이 없습니다.")

        # 통합된 DOCX 파일 저장
        docx_dir = Path(settings.BASE_DIR) / "policy_documents" / "docx" / company
        docx_dir.mkdir(parents=True, exist_ok=True)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        docx_filename = f"{company}_통합이용약관_{timestamp}.docx"
        docx_path = docx_dir / docx_filename

        doc.save(docx_path)

        logger.info(
            f"멀티 PDF → 단일 DOCX 변환 완료: {file_count}개 파일 → {docx_path}"
        )

        return {
            "success": True,
            "docx_path": str(docx_path),
            "file_count": file_count,
            "all_text": all_text,
        }

    except Exception as e:
        logger.error(f"멀티 PDF → DOCX 변환 실패: {e}")
        return {"success": False, "error": str(e)}


def process_multiple_document_upload(company, files):
    """멀티 파일 업로드 처리"""
    try:
        # 1. PDF 파일들을 임시 저장
        pdf_dir = Path(settings.BASE_DIR) / "policy_documents" / "pdf" / company
        pdf_dir.mkdir(parents=True, exist_ok=True)

        saved_pdf_files = []
        for file in files:
            if file.name.lower().endswith(".pdf"):
                filename = secure_filename(file.name)
                pdf_path = pdf_dir / filename

                with open(pdf_path, "wb") as f:
                    for chunk in file.chunks():
                        f.write(chunk)

                saved_pdf_files.append(pdf_path)

        if not saved_pdf_files:
            return {"success": False, "error": "PDF 파일이 없습니다."}

        # 2. 여러 PDF를 하나의 DOCX로 통합 변환
        conversion_result = convert_multiple_pdfs_to_single_docx(
            saved_pdf_files, company
        )

        if not conversion_result["success"]:
            return conversion_result

        docx_path = conversion_result["docx_path"]
        file_count = conversion_result["file_count"]

        # 3. 데이터베이스에 문서 정보 저장
        from datetime import datetime

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        document = PolicyDocument.objects.create(
            title=f"{company} 통합 이용약관",
            company=InsuranceCompany.objects.get(name=company),
            document_type="통합이용약관",
            description=f"{file_count}개 PDF 파일을 통합한 이용약관",
            tags=f"{company},통합,이용약관",
            file_path=docx_path,
            status="approved",
        )

        # 4. Pinecone에서 기존 해당 보험사 데이터 삭제 후 새로 업로드
        rag_service = RAGService()

        # 기존 데이터 삭제
        if rag_service.pinecone_index:
            try:
                rag_service.pinecone_index.delete(filter={"company": company})
                logger.info(f"{company} 기존 Pinecone 데이터 삭제 완료")
            except Exception as e:
                logger.warning(f"기존 데이터 삭제 실패: {e}")

        # 새로운 통합 문서 업로드
        upload_result = rag_service.upload_document_with_metadata(
            docx_path,
            company,
            "통합이용약관",
            f"{company} 통합 이용약관",
            f"{company},통합,이용약관",
        )

        return {
            "success": True,
            "message": f"{company} 통합 문서 업로드 완료: {file_count}개 파일",
            "document_id": document.id,
            "file_count": file_count,
            "pinecone_result": upload_result,
        }

    except Exception as e:
        logger.error(f"멀티 문서 업로드 실패: {e}")
        return {"success": False, "error": str(e)}


@admin_required
def admin_multiple_upload_document(request):
    """관리자 멀티 파일 업로드 페이지"""
    if request.method == "POST":
        try:
            company = request.POST.get("company")
            files = request.FILES.getlist("files")

            if not company or not files:
                return JsonResponse(
                    {"success": False, "error": "보험사와 파일을 선택해주세요."}
                )

            # 파일 처리
            result = process_multiple_document_upload(company, files)
            return JsonResponse(result)

        except Exception as e:
            return JsonResponse({"success": False, "error": str(e)})

    # GET 요청: 업로드 페이지 표시
    insurance_companies = [
        "삼성화재",
        "현대해상",
        "KB손해보험",
        "메리츠화재",
        "DB손해보험",
        "롯데손해보험",
        "하나손해보험",
        "흥국화재",
        "AXA손해보험",
        "MG손해보험",
        "캐롯손해보험",
        "한화손해보험",
    ]

    context = {"insurance_companies": insurance_companies, "title": "멀티 파일 업로드"}
    return render(request, "insurance/admin_multiple_upload.jinja.html", context)


@admin_required
def admin_document_list(request):
    """관리자 문서 목록 페이지"""
    documents = PolicyDocument.objects.all().order_by("-upload_date")
    context = {"documents": documents, "title": "문서 관리"}
    return render(request, "insurance/admin_documents.jinja.html", context)


@admin_required
def admin_pinecone_management(request):
    """관리자 Pinecone 관리 페이지"""
    rag_service = RAGService()
    stats = rag_service.get_company_document_stats()

    context = {"stats": stats, "title": "Pinecone 관리"}
    return render(request, "insurance/admin_pinecone.jinja.html", context)


# Pinecone 관리 API 뷰 함수들
@admin_required
def update_company_index(request):
    """보험사별 인덱스 업데이트"""
    if request.method == "POST":
        try:
            import json

            data = json.loads(request.body)
            company = data.get("company")

            if not company:
                return JsonResponse(
                    {"success": False, "error": "보험사 정보가 누락되었습니다."}
                )

            rag_service = RAGService()

            # 해당 보험사의 승인된 문서들만 다시 업로드
            documents = PolicyDocument.objects.filter(
                company__name=company, status="approved"
            )

            success_count = 0
            for document in documents:
                if document.file_path:
                    try:
                        result = rag_service.upload_document_with_metadata(
                            document.file_path,
                            company,
                            document.document_type,
                            document.title,
                            document.tags,
                        )
                        if result["success"]:
                            success_count += 1
                    except Exception as e:
                        logger.error(f"문서 업로드 실패: {document.title} - {e}")

            return JsonResponse(
                {
                    "success": True,
                    "message": f"{company} 인덱스 업데이트 완료: {success_count}개 문서",
                    "updated_count": success_count,
                }
            )

        except Exception as e:
            return JsonResponse({"success": False, "error": str(e)})

    return JsonResponse({"success": False, "error": "POST 요청만 지원합니다."})


@admin_required
def delete_company_data(request):
    """보험사별 데이터 삭제"""
    if request.method == "DELETE":
        try:
            import json

            data = json.loads(request.body)
            company = data.get("company")

            if not company:
                return JsonResponse(
                    {"success": False, "error": "보험사 정보가 누락되었습니다."}
                )

            rag_service = RAGService()

            # Pinecone에서 해당 보험사 데이터 삭제
            if rag_service.pinecone_index:
                # 보험사별 필터로 벡터 삭제
                rag_service.pinecone_index.delete(filter={"company": company})

            return JsonResponse(
                {"success": True, "message": f"{company} 데이터 삭제 완료"}
            )

        except Exception as e:
            return JsonResponse({"success": False, "error": str(e)})

    return JsonResponse({"success": False, "error": "DELETE 요청만 지원합니다."})


@admin_required
def update_all_index(request):
    """전체 인덱스 업데이트"""
    if request.method == "POST":
        try:
            rag_service = RAGService()

            # 기존 인덱스 초기화
            if rag_service.pinecone_index:
                rag_service.pinecone_index.delete(delete_all=True)

            # 승인된 모든 문서 다시 업로드
            documents = PolicyDocument.objects.filter(status="approved")

            success_count = 0
            for document in documents:
                if document.file_path:
                    try:
                        result = rag_service.upload_document_with_metadata(
                            document.file_path,
                            document.company.name,
                            document.document_type,
                            document.title,
                            document.tags,
                        )
                        if result["success"]:
                            success_count += 1
                    except Exception as e:
                        logger.error(f"문서 업로드 실패: {document.title} - {e}")

            return JsonResponse(
                {
                    "success": True,
                    "message": f"전체 인덱스 업데이트 완료: {success_count}개 문서",
                    "updated_count": success_count,
                }
            )

        except Exception as e:
            return JsonResponse({"success": False, "error": str(e)})

    return JsonResponse({"success": False, "error": "POST 요청만 지원합니다."})


@admin_required
def clear_all_index(request):
    """전체 인덱스 초기화"""
    if request.method == "DELETE":
        try:
            rag_service = RAGService()

            if rag_service.pinecone_index:
                rag_service.pinecone_index.delete(delete_all=True)

            return JsonResponse({"success": True, "message": "전체 인덱스 초기화 완료"})

        except Exception as e:
            return JsonResponse({"success": False, "error": str(e)})

    return JsonResponse({"success": False, "error": "DELETE 요청만 지원합니다."})


@admin_required
def get_pinecone_stats(request):
    """Pinecone 통계 조회"""
    try:
        rag_service = RAGService()
        stats = rag_service.get_company_document_stats()

        return JsonResponse({"success": True, "stats": stats})

    except Exception as e:
        return JsonResponse({"success": False, "error": str(e)})


@admin_required
def delete_document_api(request, document_id):
    """문서 삭제 API"""
    if request.method == "DELETE":
        try:
            document = PolicyDocument.objects.get(id=document_id)

            # Pinecone에서 해당 문서 벡터 삭제
            rag_service = RAGService()
            if rag_service.pinecone_index:
                rag_service.pinecone_index.delete(
                    filter={
                        "source": document.file_path.split("/")[-1],
                        "company": document.company.name,
                    }
                )

            # 데이터베이스에서 문서 삭제
            document.delete()

            return JsonResponse({"success": True, "message": "문서 삭제 완료"})

        except PolicyDocument.DoesNotExist:
            return JsonResponse({"success": False, "error": "문서를 찾을 수 없습니다."})
        except Exception as e:
            return JsonResponse({"success": False, "error": str(e)})

    return JsonResponse({"success": False, "error": "DELETE 요청만 지원합니다."})


def secure_filename(filename):
    """안전한 파일명 생성 (Django 내장 함수 사용)"""
    # 파일명에서 특수문자 제거
    filename = re.sub(r"[^\w\s-]", "", filename)
    # 공백을 언더스코어로 변경
    filename = re.sub(r"[-\s]+", "_", filename)
    return filename


def search_documents_api(request):
    """문서 검색 API"""
    if request.method == "GET":
        try:
            query = request.GET.get("query", "")
            company = request.GET.get("company", "")
            top_k = int(request.GET.get("top_k", 5))

            if not query:
                return JsonResponse({"success": False, "error": "검색어가 필요합니다."})

            rag_service = RAGService()

            if company:
                results = rag_service.search_documents_by_company(query, company, top_k)
            else:
                results = rag_service.search_documents(query, top_k)

            return JsonResponse(
                {
                    "success": True,
                    "results": results,
                    "query": query,
                    "company": company,
                    "count": len(results),
                }
            )

        except Exception as e:
            return JsonResponse({"success": False, "error": str(e)})

    return JsonResponse({"success": False, "error": "GET 요청만 지원합니다."})


def convert_pdf_to_docx_optimized(pdf_path, docx_path):
    """PDF를 DOCX로 변환 (용량 최적화 버전)"""
    try:
        # 방법 1: PyPDF2 사용
        text = extract_text_with_pypdf2_optimized(pdf_path)

        # 방법 1이 실패하면 방법 2 시도
        if not text.strip():
            text = extract_text_with_alternative_method_optimized(pdf_path)

        if not text.strip():
            raise Exception("PDF에서 텍스트를 추출할 수 없습니다.")

        # DOCX 생성 (용량 최적화)
        doc = Document()

        # 문서 속성 최적화
        doc.core_properties.title = "자동차보험 약관"
        doc.core_properties.author = "보험 추천 시스템"
        doc.core_properties.subject = "자동차보험 약관"

        # 텍스트를 문단으로 분할하여 추가 (용량 최적화)
        paragraphs = text.split("\n")
        for para in paragraphs:
            if para.strip():
                # 특수 문자 제거 후 문단 추가
                clean_para = clean_text_for_docx_optimized(para)
                if clean_para.strip():
                    doc.add_paragraph(clean_para)

        doc.save(docx_path)
        logger.info(f"PDF → DOCX 변환 완료 (최적화): {pdf_path} → {docx_path}")

    except Exception as e:
        logger.error(f"PDF → DOCX 변환 실패: {e}")
        raise e


def extract_text_with_pypdf2_optimized(pdf_path):
    """PyPDF2를 사용한 텍스트 추출 (용량 최적화)"""
    try:
        import PyPDF2

        with open(pdf_path, "rb") as file:
            pdf_reader = PyPDF2.PdfReader(file)

            text = ""
            for page_num in range(len(pdf_reader.pages)):
                try:
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()

                    if page_text:
                        # 한글 텍스트 정리 (용량 최적화)
                        cleaned_text = clean_korean_text_optimized(page_text)
                        if cleaned_text.strip():
                            text += cleaned_text + "\n\n"

                except Exception as e:
                    logger.warning(f"페이지 {page_num + 1} 텍스트 추출 실패: {e}")
                    continue

            return text.strip()

    except Exception as e:
        logger.error(f"PyPDF2 텍스트 추출 실패: {e}")
        return ""


def clean_korean_text_optimized(text):
    """한글 텍스트 정리 (용량 최적화)"""
    import re

    # 기본 정리
    text = text.replace("\x00", "")  # NULL 바이트 제거

    # 불필요한 공백 제거
    text = re.sub(r"\s+", " ", text)

    # 빈 줄 정리
    text = re.sub(r"\n\s*\n", "\n", text)

    # 한글 문자만 유지 (용량 최적화)
    text = re.sub(r"[^\uAC00-\uD7AF\w\s\.\,\;\:\!\?\(\)\[\]\-\_]", "", text)

    return text.strip()


def extract_text_with_alternative_method_optimized(pdf_path):
    """대체 방법을 사용한 텍스트 추출 (용량 최적화)"""
    try:
        # 방법 1: pdfplumber 시도
        try:
            import pdfplumber

            with pdfplumber.open(pdf_path) as pdf:
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        cleaned_text = clean_korean_text_optimized(page_text)
                        if cleaned_text.strip():
                            text += cleaned_text + "\n\n"

                if text.strip():
                    return clean_text_for_docx_optimized(text)
        except ImportError:
            logger.info("pdfplumber가 설치되지 않음, 다른 방법 시도")
        except Exception as e:
            logger.warning(f"pdfplumber 추출 실패: {e}")

        # 방법 2: PyMuPDF (fitz) 시도
        try:
            import fitz

            doc = fitz.open(pdf_path)
            text = ""
            for page in doc:
                page_text = page.get_text()
                if page_text:
                    cleaned_text = clean_korean_text_optimized(page_text)
                    if cleaned_text.strip():
                        text += cleaned_text + "\n\n"
            doc.close()

            if text.strip():
                return clean_text_for_docx_optimized(text)
        except ImportError:
            logger.info("PyMuPDF가 설치되지 않음, 다른 방법 시도")
        except Exception as e:
            logger.warning(f"PyMuPDF 추출 실패: {e}")

        return ""

    except Exception as e:
        logger.warning(f"대체 방법 추출 실패: {e}")
        return ""


def clean_text_for_docx_optimized(text):
    """DOCX 호환을 위한 텍스트 정리 (용량 최적화)"""
    import re

    # NULL 바이트 제거
    text = text.replace("\x00", "")

    # 제어 문자 제거
    text = "".join(char for char in text if ord(char) >= 32 or char in "\n\r\t")

    # 연속된 공백 정리
    text = re.sub(r"\s+", " ", text)

    # 빈 줄 정리
    text = re.sub(r"\n\s*\n", "\n", text)

    return text.strip()


def convert_all_pdfs_to_docx():
    """모든 PDF 파일을 DOCX로 변환"""
    try:
        from pathlib import Path
        from datetime import datetime

        pdf_base_dir = Path(settings.BASE_DIR) / "policy_documents" / "pdf"
        docx_base_dir = Path(settings.BASE_DIR) / "policy_documents" / "docx"

        if not pdf_base_dir.exists():
            return {"success": False, "error": "PDF 폴더가 존재하지 않습니다."}

        conversion_results = []
        total_files = 0
        success_count = 0

        # 각 보험사 폴더 순회
        for company_dir in pdf_base_dir.iterdir():
            if company_dir.is_dir():
                company_name = company_dir.name
                docx_company_dir = docx_base_dir / company_name
                docx_company_dir.mkdir(parents=True, exist_ok=True)

                # 기존 DOCX 파일 삭제
                for docx_file in docx_company_dir.glob("*.docx"):
                    try:
                        docx_file.unlink()
                        logger.info(f"기존 DOCX 파일 삭제: {docx_file}")
                    except Exception as e:
                        logger.warning(f"기존 DOCX 파일 삭제 실패: {docx_file} - {e}")

                # PDF 파일들 변환
                for pdf_file in company_dir.glob("*.pdf"):
                    total_files += 1
                    try:
                        # DOCX 파일명 생성
                        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                        docx_filename = f"{pdf_file.stem}_{timestamp}.docx"
                        docx_path = docx_company_dir / docx_filename

                        # PDF → DOCX 변환
                        convert_pdf_to_docx_optimized(pdf_file, docx_path)

                        conversion_results.append(
                            {
                                "company": company_name,
                                "pdf_file": pdf_file.name,
                                "docx_file": docx_filename,
                                "status": "success",
                            }
                        )
                        success_count += 1

                    except Exception as e:
                        logger.error(f"PDF 변환 실패: {pdf_file} - {e}")
                        conversion_results.append(
                            {
                                "company": company_name,
                                "pdf_file": pdf_file.name,
                                "status": "failed",
                                "error": str(e),
                            }
                        )

        return {
            "success": True,
            "message": f"PDF 변환 완료: {success_count}/{total_files}개 파일",
            "total_files": total_files,
            "success_count": success_count,
            "results": conversion_results,
        }

    except Exception as e:
        logger.error(f"전체 PDF 변환 실패: {e}")
        return {"success": False, "error": str(e)}


def process_hanwha_insurance_upload():
    """한화손해보험 PDF 변환 및 Pinecone 업데이트"""
    try:
        from pathlib import Path
        from datetime import datetime

        # 파일 경로 설정
        pdf_path = (
            Path(settings.BASE_DIR)
            / "policy_documents"
            / "pdf"
            / "한화손해보험"
            / "(취합)한화손해보험_개인용(공동물건)자동차보험_약관.pdf"
        )
        docx_dir = (
            Path(settings.BASE_DIR) / "policy_documents" / "docx" / "한화손해보험"
        )
        docx_dir.mkdir(parents=True, exist_ok=True)

        # 기존 DOCX 파일 삭제
        for docx_file in docx_dir.glob("*.docx"):
            try:
                docx_file.unlink()
                logger.info(f"기존 DOCX 파일 삭제: {docx_file}")
            except Exception as e:
                logger.warning(f"기존 DOCX 파일 삭제 실패: {docx_file} - {e}")

        # DOCX 파일명 생성
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        docx_filename = f"한화손해보험_자동차보험_약관_{timestamp}.docx"
        docx_path = docx_dir / docx_filename

        # PDF 파일 존재 확인
        if not pdf_path.exists():
            return {
                "success": False,
                "error": "한화손해보험 PDF 파일이 존재하지 않습니다.",
            }

        # PDF → DOCX 변환
        convert_pdf_to_docx_optimized(pdf_path, docx_path)

        # Pinecone에서 기존 한화손해보험 데이터 삭제
        rag_service = RAGService()
        if rag_service.pinecone_index:
            try:
                rag_service.pinecone_index.delete(filter={"company": "한화손해보험"})
                logger.info("한화손해보험 기존 Pinecone 데이터 삭제 완료")
            except Exception as e:
                logger.warning(f"기존 데이터 삭제 실패: {e}")

        # 새로운 DOCX 파일을 Pinecone에 업로드
        upload_result = rag_service.upload_document_with_metadata(
            docx_path,
            "한화손해보험",
            "이용약관",
            "한화손해보험 자동차보험 약관",
            "한화손해보험,자동차보험,약관",
        )

        return {
            "success": True,
            "message": "한화손해보험 PDF 변환 및 Pinecone 업데이트 완료",
            "pdf_file": pdf_path.name,
            "docx_file": docx_filename,
            "pinecone_result": upload_result,
        }

    except Exception as e:
        logger.error(f"한화손해보험 처리 실패: {e}")
        return {"success": False, "error": str(e)}


@admin_required
def admin_convert_all_pdfs(request):
    """모든 PDF 파일 변환 관리 페이지"""
    if request.method == "POST":
        try:
            # 모든 PDF → DOCX 변환
            result = convert_all_pdfs_to_docx()
            return JsonResponse(result)

        except Exception as e:
            return JsonResponse({"success": False, "error": str(e)})

    # GET 요청: 변환 페이지 표시
    context = {
        "title": "전체 PDF 변환",
        "description": "모든 PDF 파일을 DOCX로 변환합니다.",
    }

    return render(request, "insurance/admin_convert_all.jinja.html", context)


@admin_required
def admin_hanwha_insurance_convert(request):
    """한화손해보험 PDF 변환 전용 페이지"""
    if request.method == "POST":
        try:
            # 한화손해보험 PDF 변환 및 Pinecone 업데이트
            result = process_hanwha_insurance_upload()
            return JsonResponse(result)

        except Exception as e:
            return JsonResponse({"success": False, "error": str(e)})

    # GET 요청: 변환 페이지 표시
    context = {
        "title": "한화손해보험 PDF 변환",
        "company": "한화손해보험",
        "pdf_filename": "(취합)한화손해보험_개인용(공동물건)자동차보험_약관.pdf",
        "pdf_size": "57MB",
    }

    return render(request, "insurance/admin_hanwha_convert.jinja.html", context)
