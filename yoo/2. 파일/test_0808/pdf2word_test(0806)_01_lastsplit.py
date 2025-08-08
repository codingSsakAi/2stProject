import fitz
import pdfplumber
from PIL import Image
import pytesseract
import io
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from tqdm import tqdm
import re
import os
import logging
import warnings

warnings.filterwarnings("ignore", message="Cannot set gray non-stroke color")

# ======================
# 유틸 함수
# ======================
def clean_invalid_chars(text):
    if not text:
        return ""
    text = text.replace('\x00', '')
    text = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]', '', text)
    text = re.sub(r'[^\u0000-\uFFFF]', '', text)
    return text.strip()

def merge_lines(lines):
    merged, buffer = [], ""
    for line in lines:
        line = line.strip()
        if not line:
            if buffer:
                merged.append(buffer.strip())
                buffer = ""
        else:
            if buffer.endswith('-'):
                buffer = buffer[:-1] + line
            elif buffer and not buffer.endswith(('.', '!', '?', ':')):
                buffer += " " + line
            else:
                if buffer:
                    merged.append(buffer.strip())
                buffer = line
    if buffer:
        merged.append(buffer.strip())
    return merged

def setup_logger(log_file):
    logger = logging.getLogger(__name__)
    logger.setLevel(logging.INFO)
    fh = logging.FileHandler(log_file, encoding="utf-8")
    formatter = logging.Formatter('%(asctime)s - %(message)s')
    fh.setFormatter(formatter)
    if logger.hasHandlers():
        logger.handlers.clear()
    logger.addHandler(fh)
    return logger

# ======================
# 표 병합
# ======================
def merge_table_cells(table, table_data):
    rows = len(table_data)
    cols = len(table_data[0])
    # 세로 병합
    for col in range(cols):
        merge_start = 0
        prev_text = table.cell(0, col).text.strip()
        for row in range(1, rows):
            curr_text = table.cell(row, col).text.strip()
            if curr_text == "" or curr_text == prev_text:
                table.cell(merge_start, col).merge(table.cell(row, col))
            else:
                prev_text = curr_text
                merge_start = row
    # 가로 병합
    for row in range(rows):
        merge_start = 0
        prev_text = table.cell(row, 0).text.strip()
        for col in range(1, cols):
            curr_text = table.cell(row, col).text.strip()
            if curr_text == prev_text and curr_text != "":
                table.cell(row, merge_start).merge(table.cell(row, col))
            else:
                prev_text = curr_text
                merge_start = col

# ======================
# 세로 텍스트 (1글자 기준)
# ======================
def detect_vertical_text(spans):
    verticals, horizontals = [], []
    for span in spans:
        text = span["text"].strip()
        if not text:
            continue
        x0, y0 = span["bbox"][0], span["bbox"][1]
        direction = span.get("dir", (1, 0))
        if (direction[0] == 0 and direction[1] != 0) or len(text) == 1:
            verticals.append((x0, y0, text))
        else:
            horizontals.append((x0, y0, text))
    return verticals, horizontals

# ======================
# PDF 추출
# ======================
def extract_page_contents(pdf_path, logger, max_pages=None):
    pdf_doc = fitz.open(pdf_path)
    total_pages = len(pdf_doc) if max_pages is None else min(len(pdf_doc), max_pages)
    results = []

    with pdfplumber.open(pdf_path) as pdf_plumber_doc:
        for page_idx in tqdm(range(total_pages), desc="PDF 페이지 처리", unit="page"):
            page = pdf_doc[page_idx]
            plumber_page = pdf_plumber_doc.pages[page_idx]
            page_content = {"items": [], "verticals": []}

            # 텍스트 추출
            spans = []
            for block in page.get_text("dict")["blocks"]:
                if "lines" in block:
                    for line in block["lines"]:
                        spans.extend(line["spans"])

            vertical_spans, horizontal_spans = detect_vertical_text(spans)

            # 세로 텍스트
            if vertical_spans:
                vertical_spans.sort(key=lambda t: (round(t[0], 1), t[1]))
                grouped = {}
                for x, y, t in vertical_spans:
                    grouped.setdefault(round(x, 1), []).append((y, t))
                for _, group in grouped.items():
                    group.sort(key=lambda g: g[0])
                    merged = "".join([g[1] for g in group])
                    merged = re.sub(r"[•\d]", "", merged)  # 숫자/• 제거
                    page_content["verticals"].append(clean_invalid_chars(merged))

            # 본문 텍스트 y->x 병합
            horizontal_spans.sort(key=lambda t: (round(t[1], 1), t[0]))
            y_grouped = {}
            for x, y, text in horizontal_spans:
                y_grouped.setdefault(round(y, 1), []).append((x, text))
            for y, texts in sorted(y_grouped.items()):
                line = " ".join([txt for _, txt in sorted(texts, key=lambda v: v[0])])
                page_content["items"].append({"type": "text", "y": y, "content": clean_invalid_chars(line)})

            # 표 추출 및 좌표
            tables = plumber_page.extract_tables()
            for table in tables:
                if any(any(cell and cell.strip() for cell in row) for row in table):
                    table_rows = [[cell.strip() if cell else "" for cell in row] for row in table]
                    words = plumber_page.extract_words()
                    min_y = min((w["top"] for w in words), default=99999)
                    page_content["items"].append({"type": "table", "y": min_y, "content": table_rows})

            # 본문-표 중복 제거
            all_table_texts = {cell for item in page_content["items"] if item["type"]=="table"
                               for row in item["content"] for cell in row if cell}
            filtered_items = []
            for item in page_content["items"]:
                if item["type"] == "text":
                    if not any(tbl_text in item["content"] for tbl_text in all_table_texts):
                        filtered_items.append(item)
                else:
                    filtered_items.append(item)
            page_content["items"] = filtered_items

            results.append(page_content)
            logger.info(f"{page_idx+1}p: 텍스트 {sum(1 for i in page_content['items'] if i['type']=='text')}, 표 {sum(1 for i in page_content['items'] if i['type']=='table')}, 세로 {len(page_content['verticals'])}")
    return results

# ======================
# PDF → Word 변환
# ======================
def pdf_to_word(pdf_path, word_path, max_pages=None):
    filename = os.path.splitext(os.path.basename(pdf_path))[0]
    log_file = f"{filename}.log"
    logger = setup_logger(log_file)
    print(f"[INFO] '{filename}' 처리 시작")

    results = extract_page_contents(pdf_path, logger, max_pages=max_pages)
    doc = Document()

    for idx, page in enumerate(results, start=1):
        sorted_items = sorted(page["items"], key=lambda i: i["y"])
        for item in sorted_items:
            if item["type"] == "text":
                p = doc.add_paragraph(item["content"])
                run = p.runs[0]
                run.font.name = '맑은 고딕'
                run.font.size = Pt(11)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            elif item["type"] == "table":
                doc.add_paragraph()
                table_data = item["content"]
                max_cols = max(len(row) for row in table_data)
                table = doc.add_table(rows=len(table_data), cols=max_cols)
                table.style = 'Table Grid'
                for i, row in enumerate(table_data):
                    for j in range(max_cols):
                        table.cell(i, j).text = row[j] if j < len(row) else ""
                merge_table_cells(table, table_data)
                doc.add_paragraph()

        if page["verticals"]:
            doc.add_paragraph("[세로 텍스트]")
            doc.add_paragraph(" ".join(page["verticals"]))
        doc.add_paragraph(f"--- {idx} 페이지 ---").runs[0].font.size = Pt(8)

    doc.save(word_path)
    print(f"[INFO] Word 저장 완료: {word_path}")

# === 실행 ===
if __name__ == "__main__":
    pdf_to_word(
        r"C:\Users\Admin\Desktop\2차 프로젝트\2.코드\dataset_pdf\롯데손해보험_원본\롯데손해보험\롯데.pdf",
        r"C:\Users\Admin\Desktop\2차 프로젝트\2.코드\dataset_pdf\롯데손해보험_원본\롯데손해보험\롯데_test(0808).docx",
        max_pages=10
    )
