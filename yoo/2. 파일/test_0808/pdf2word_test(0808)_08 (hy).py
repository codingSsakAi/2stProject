import fitz  # PyMuPDF
import re
import os
import warnings
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from hanspell import spell_checker
from konlpy.tag import Okt

warnings.filterwarnings("ignore", category=UserWarning)

def clean_invalid_chars(text):
    if not text:
        return ""
    text = text.replace('\x00', '')
    text = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]', '', text)  # 제어문자 제거
    text = re.sub(r'[^\u0000-\uFFFF]', '', text)  # 유니코드 외 문자 제거
    return text.strip()

def correct_text(text):
    try:
        result = spell_checker.check(text)
        return result.checked
    except:
        return text  # 오류 시 원문 반환

def merge_lines(lines, y_threshold=2.0):
    merged = []
    buffer = ""
    last_y = None
    for line in lines:
        y, content = line
        content = content.strip()
        if not content:
            continue
        if last_y is not None and abs(y - last_y) > y_threshold:
            merged.append(buffer.strip())
            buffer = content
        else:
            if buffer:
                buffer += ' ' + content
            else:
                buffer = content
        last_y = y
    if buffer:
        merged.append(buffer.strip())
    return merged

def extract_text_from_pdf(pdf_path, max_pages=10):
    doc = fitz.open(pdf_path)
    results = []
    vertical_texts = []
    header_footer_candidates = {}

    for page_index in range(min(len(doc), max_pages)):
        page = doc[page_index]
        words = page.get_text("words")
        lines = {}
        verticals = []

        for w in words:
            x0, y0, x1, y1, word, block_no, line_no, word_no = w
            word = clean_invalid_chars(word)
            if len(word) == 1 and re.fullmatch(r'[가-힣]', word):
                verticals.append((x0, y0, word))
                continue
            y_key = round(y0, 1)
            lines.setdefault(y_key, []).append((x0, word))

        vertical_texts.extend(verticals)

        # 병합 및 정렬
        line_items = []
        for y, texts in lines.items():
            sorted_line = " ".join(t for _, t in sorted(texts, key=lambda x: x[0]))
            line_items.append((y, sorted_line))

        merged = merge_lines(sorted(line_items, key=lambda x: x[0]))

        # 헤더/풋터 후보 수집
        for m in merged:
            header_footer_candidates[m] = header_footer_candidates.get(m, 0) + 1

        # 맞춤법 보정
        corrected = [correct_text(line) for line in merged if line]
        results.append(corrected)

    # 머리말/꼬리말 제거
    total_pages = len(results)
    filtered_results = []
    for page_lines in results:
        filtered = [line for line in page_lines if header_footer_candidates.get(line, 0) < total_pages * 0.6]
        filtered_results.append(filtered)

    # 세로 텍스트 마지막에 병합
    if vertical_texts:
        grouped = {}
        for x, y, char in vertical_texts:
            grouped.setdefault(round(x, 1), []).append((y, char))
        vertical_lines = []
        for _, chars in grouped.items():
            chars.sort(key=lambda x: x[0])
            vline = "".join(c for _, c in chars)
            if len(vline.strip()) > 2:
                vertical_lines.append(vline)
        filtered_results.append(vertical_lines)

    return filtered_results

def save_to_docx(pages, output_path):
    doc = Document()
    for idx, lines in enumerate(pages, start=1):
        for line in lines:
            p = doc.add_paragraph(line)
            run = p.runs[0]
            run.font.name = '맑은 고딕'
            run.font.size = Pt(10.5)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        doc.add_paragraph(f"--- {idx}페이지 ---")
    doc.save(output_path)

def save_to_md(pages, output_path):
    with open(output_path, "w", encoding="utf-8") as f:
        for idx, lines in enumerate(pages, start=1):
            for line in lines:
                f.write(f"{line}\n")
            f.write(f"\n--- {idx}페이지 ---\n\n")

if __name__ == "__main__":
    input_pdf = r"C:/Users/Admin/Desktop/2차 프로젝트/2.코드/dataset_pdf/롯데손해보험_원본/롯데손해보험/롯데.pdf"
    output_docx = input_pdf.replace(".pdf", "_형태소.docx")
    output_md = input_pdf.replace(".pdf", "_형태소.md")

    print("PDF 분석 중... 최대 10페이지")
    result_pages = extract_text_from_pdf(input_pdf, max_pages=10)

    print("DOCX 저장 중...")
    save_to_docx(result_pages, output_docx)

    print("MD 저장 중...")
    save_to_md(result_pages, output_md)

    print("✅ 완료:", output_docx)
    print("✅ 완료:", output_md)
